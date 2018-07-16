import gevent.monkey
gevent.monkey.patch_all()

import os
from flask import Flask, request, jsonify, abort, g, render_template
from flask_cors import CORS
from flask_httpauth import HTTPBasicAuth
from sqlalchemy.exc import IntegrityError
from sqlalchemy import Time, cast
from gevent.pywsgi import WSGIServer
from flask_mail import Mail, Message
from flask import send_file
from app import db, create_app,socketio
from app.models.Tasks import Tasks
from app.models.Notification import Notification
from app.models.User import User
from app.models.Lectures import Lectures
from app.models.Batch import Batch
from app.models.Role import Role
from app.models.Department import Department
from app.models.userrole import UserRoles
from app.models.deptmap import DeptMap
from app.models.Application import Application
from app.models.SocketClient import SocketClient
from app.models.FacultyAttendance import FacultyAttendance
from format import m, all_rooms, all_batches,fingerPrintCode, weekday_map_dict
from datetime import datetime
from datetime import timedelta
from dotenv import load_dotenv
from scripts.timetable import timetable
from scripts.approvals import short_leave
from flask_socketio import join_room, leave_room
import gevent
import pgpubsub
import json
import glob


auth = HTTPBasicAuth()

app = create_app()

mail = Mail(app)


# decorator for verifying role
def roles_required(*role_names):
    """| This decorator ensures that the current user is logged in,
    | and has *all* of the specified roles (AND operation).
    """
    def wrapper(view_function):

        def decorator(*args, **kwargs):
            user = g.user
            if not user:
                render_template('result.html', msg="not authorized ")
            # User must have the required roles
            if not user.has_roles(*role_names):
                # Redirect to the unauthorized page
                return render_template('result.html', msg="not for this role")
            # It's OK to call the view
            return view_function(*args, **kwargs)
        decorator.__name__ = view_function.__name__
        return decorator
    return wrapper


# load dotenv in the base root
dotenv_path = os.path.join(os.path.dirname(__file__), '.env')
load_dotenv(dotenv_path)

app.config['SQLALCHEMY_DATABASE_URI'] = os.environ['DATABASE_URL']
db.init_app(app)


@app.route('/users', methods=['POST'])
def new_user():
    email = request.json.get('email')
    password = request.json.get('password')
    name = request.json.get('name')
    profile_url = request.json.get('profile_url')
    profile_photo_url = request.json.get('profile_photo_url')
    designation = request.json.get('designation')
    department = None or request.json.get('department').lower()
    number = request.json.get('number')
    role = request.json.get('role')
    if department:
        department = department.lower()

    _role = Role.query.filter_by(name=role).first()
    _department = Department.query.filter_by(department=department).first()

    if email is None or password is None:
        abort(400)  # missing arguments
    if User.query.filter_by(email=email).first() is not None:
        abort(400)  # existing user

    user = User(email, profile_photo_url, number, profile_url,
                name, None, designation, None)
    if not _role:
        _role = Role(name=role)
    if not _department:
        _department = Department(name=department)
        db.session.add(_department)
    if _role is "Hod":

        _department.hod_user_id = user.id
        db.session.add(_department)
    user.department_id = [_department]
    user.roles = [_role]
    user.set_password(password)
    db.session.add(user)
    _date1 = datetime.now().date()
    _date2 = _date1 + timedelta(days=1)
    _t1 = Tasks("bvcoe.erp@gmail.com", email, "vansifysecret", _date1,
                'https://vignette.wikia.nocookie.net/monstergirlencyclopedia/images/d/d0/Discord_logo.png/revision/latest?cb=20160622232221', _date1)
    _t2 = Tasks("bvcoe.erp@gmail.com", email, "vansifysecret", _date2,
                'https://vignette.wikia.nocookie.net/monstergirlencyclopedia/images/d/d0/Discord_logo.png/revision/latest?cb=20160622232221', _date2)
    db.session.add(_t1)
    db.session.add(_t2)
    db.session.commit()
    return jsonify({'username': user.username})


@app.route('/getusers')
@auth.login_required
@roles_required('Tte', 'Admin')
def get_users():
    _users = User.query.all()
    return jsonify(
        [user.email for user in _users]
    )


@app.route('/resource')
@auth.login_required
def get_resource():
    return jsonify({'data': 'Hello %s' % g.user.username})


@app.route('/token', methods=['POST', 'GET'])
@auth.login_required
def get_auth_token():
    role = request.json.get('role')
    if not g.user.has_roles(role):
        print(role)
        abort(401)
    token = g.user.generate_auth_token()
    return jsonify({
        'token': token.decode('ascii'),
        "email": g.user.username,
        "name": g.user.name,
        "profile_photo_url": g.user.profile_photo_url,
        "role": role
    })


@auth.verify_password
def verify_password(username_or_token, password):
    # first try to authenticate by token
    user = User.verify_auth_token(username_or_token.encode("utf-8"))
    if not user:
        # try to authenticate with username/password
        user = User.query.filter_by(email=username_or_token).first()
        if not user or not user.check_password(password):
            return False
    g.user = user
    return True


@app.route("/notif")
@auth.login_required
def notif():
    s = g.user.username
    r = request.args['r']
    m = request.args['m']
    prof_url = g.user.profile_photo_url
    n = Notification(s, r, m, prof_url)
    db.session.add(n)
    try:
        db.session.commit()
    except IntegrityError:
        pass

    new_notifs = Notification.query.filter_by(r_email=r).all()

    return jsonify(
        [notif.id for notif in new_notifs]
    )


@app.route("/notifs")
@auth.login_required
def notifs():
    r = g.user.username
    notifs = Notification.query.filter_by(r_email=r).all()
    return jsonify(
        [notif.json() for notif in notifs]
    )


@app.route("/task")
@auth.login_required
def task():
    s = g.user.username
    r = g.user.username
    t = request.args['t']
    tcd = request.args['tcd']
    p = g.user.profile_photo_url
    _t = Tasks(s, r, t, tcd, p)
    db.session.add(_t)
    try:
        db.session.commit()
        return jsonify(Tasks.query.filter_by(r_email=r).order_by(Tasks.id.desc()).first().json())
    except IntegrityError:
        return "Integrity Error"


@app.route("/taskby")
@auth.login_required
def taskby():
    s = request.args['s']
    r = g.user.email
    t = request.args['t']
    tcd = request.args['tcd']
    prof_photo_url = User.query.with_entities(User.profile_photo_url)\
                         .filter(User.email == s)\
                         .first()

    _t = Tasks(s, r, t, tcd, prof_photo_url)
    db.session.add(_t)
    try:
        db.session.commit()
        return jsonify(Tasks.query.filter_by(r_email=r).order_by(Tasks.id.desc()).first().json())
    except IntegrityError:
        return "Integrity Error"


@app.route("/tasks")
@auth.login_required
def tasks():
    r = g.user.email
    d = request.args['d']
    _tasks = Tasks.query.filter(Tasks.r_email == g.user.email).filter(
        Tasks.to_complete_date == d).all()
    _tasks = [task.json() for task in _tasks]
    weekday = datetime.strptime(str(d), '%Y-%m-%d').weekday()
    day = weekday_map_dict[weekday]
    _lectures = Lectures.query.filter(Lectures.t_email == g.user.email).filter(Lectures.day == day).order_by(Lectures.s).all()
    _lectures = [lec.to_task(d) for lec in _lectures]
    return jsonify(_tasks + _lectures)


@app.route("/contacts")
@auth.login_required
def return_contacts():
    _users = User.query.all()
    return jsonify([user.json() for user in _users])


def time(x):
    return datetime.strptime(str(x), '%I:%M%p').replace(second=0, microsecond=0).time()


@app.route("/getlist")
@auth.login_required
def avail():
    s_time = request.args['s']
    day = request.args['day'].lower()

    start = time(s_time)

    _busy_rooms = Lectures.query.filter(day == Lectures.day)\
                                .filter((cast(Lectures.e, Time) > start) &
                                        (cast(Lectures.s, Time) < start))\
                                .all()
    _free = []
    b_info = {}
    possibilities = [["full"], ["p1", "p2", "p3"], ["t1", "t2"]]
    for room in _busy_rooms:
        if room.batch_name not in b_info.keys():
            b_info[room.batch_name] = []
        b_info[room.batch_name].append(room.batch_part)

    b_names = [room.batch_name for room in _busy_rooms]
    all_teachers = [user.username for user in User.query.all()]
    b_teachers = [room.t_email for room in _busy_rooms]
    print(_busy_rooms)
    _free_teachers = list(set(all_teachers).difference(b_teachers))
    _free_batches = list(set(all_batches.keys()).difference(b_names))

    print(b_info)
    for name in b_info.keys():
        if b_info[name] == possibilities[0]:
            _free.append({"label":None, "value": None, "children":None})
        elif all(x in possibilities[1] for x in b_info[name]):
            put = list(set(possibilities[1]) - set(b_info[name]))
            print(put)
            child = []
            for p in put:
                child.append({"label":p, "value":p})
            _free.append({"label":name, "value":name, "children": child})
            print(_free)
        elif all(x in possibilities[2] for x in b_info[name]):
            put = list(set(possibilities[2]) - set(b_info[name]))
            child = []
            for p in put:
                child.append({"label":p, "value":p})
            _free.append({"label":name, "value":name, "children": child})
    _free_batch_opt = [
        {"label": "full", "value": "full"},
        {"label": "p1", "value": "p1"},
        {"label": "p2", "value": "p2"},
        {"label": "p3", "value": "p3"},
        {"label": "t1", "value": "t1"},
        {"label": "t2", "value": "t2"}]

    for batch in _free_batches:
        _free.append({"label": batch, "value": batch, "children": _free_batch_opt})

    _free_teachers = [{"label": teacher, "value": teacher}
                      for teacher in _free_teachers]
    data = {
        "free_teachers": _free_teachers,
        "free_batches": _free
    }
    return jsonify(data)


@app.route("/addlec", methods=['POST'])
@auth.login_required
@roles_required('Tte')
def add_lec_info():
    # TODO: Add Session field to the lectures table
    t_email = request.json.get('t_email')[0]
    t_name = User.query.with_entities(User.name).filter(
        t_email == User.username).first(),
    lec_name = request.json.get('lec_name').lower()
    room_id = request.json.get('room_id').lower()
    batch = request.json.get('batch')
    batch_name = batch[0].lower()
    batch_part = batch[1].lower()
    day = request.json.get('day').lower()

    # Read start and end time in str format
    start_time = request.json.get('s')
    end_time = request.json.get('e')

    # convert start_time,end_time format from str to datetime
    start = time(start_time)
    end = time(end_time)

    _check = Lectures.query.filter(Lectures.room_id == room_id)\
                           .filter(day == Lectures.day)\
                           .filter((cast(Lectures.e, Time) >= start) & (cast(Lectures.s, Time) <= end))\
                           .first()

    if _check is not None:
        abort(400)

    date = datetime.now().date()
    start = datetime.combine(date, start)
    end = datetime.combine(date, end)

    _lec_info = Lectures(t_name, t_email, lec_name,
                         room_id, start, end, day, batch_name, batch_part)

    db.session.add(_lec_info)

    try:
        db.session.commit()
        return "Success"
    except IntegrityError:
        return "Integrity Error"


@app.route("/dellec", methods=['POST'])
@auth.login_required
@roles_required('Tte')
def del_lec():
    _id = request.args['id']
    Lectures.query.filter(Lectures.id == _id).delete()
    try:
        db.session.commit()
        return "Success"
    except IntegrityError:
        return "Integrity Error"


@app.route("/get-room-info")
@auth.login_required
def get_room_info():

    day = request.args['day'].lower()

    s_time = request.args['s']
    start = time(s_time)
    # End Time will be neccessary in future versions so that we can book the room in particular time slot
    # e_time = request.json.get('end_time')
    # end = time(e_time)

    _busy_rooms = Lectures.query.with_entities(Lectures.room_id)\
                                .filter(day == Lectures.day)\
                                .filter((cast(Lectures.e, Time) >= start) & (cast(Lectures.s, Time) <= start))\
                                .all()

    _busy_rooms = [room[0] for room in _busy_rooms]

    return jsonify({
        "busyrooms": _busy_rooms,
        "freerooms": list(set(all_rooms).difference(_busy_rooms))
    })


@app.route("/getsubs", methods=['GET'])
@auth.login_required
def getsubs():
    subject = request.args['sub']
    return jsonify(all_batches[subject])


@app.route("/getlecs", methods=['GET'])
@auth.login_required
def get_lecs():
    type = request.args['type'].lower()

    if type == 'room':
        room_id = request.args['id'].lower()
        _lecs = Lectures.query.filter(Lectures.room_id == room_id).all()
        return jsonify([lec.json() for lec in _lecs])

    elif type == 'faculty':
        email = request.args['id']
        _lecs = Lectures.query.filter(Lectures.t_email == email).all()
        return jsonify([lec.json() for lec in _lecs])

    elif type == 'batch':
        batch = request.args['id'].lower()
        _lecs = Lectures.query.filter(Lectures.batch_name == batch).all()
        return jsonify([lec.json() for lec in _lecs])

    else:
        print("invalid type")
        return abort(400)


@app.route("/downloadtt", methods=['GET'])
def download_tt():

    # filelist = glob.glob('{}/timetable/*'.format(os.environ['FILE_PATH']))
    # for f in filelist:
    #     os.remove(f)

    type_ = request.args['type'].lower()
    ext = request.args['ext'].lower()

    if type_ == 'room' or type_ is None:
        room = request.args['id'].lower()
        _lecs = Lectures.query.filter(Lectures.room_id == room).all()
        f = timetable.room_timetable([lec.json() for lec in _lecs], room)
        if ext == 'docx':
            try:
                return send_file(f, as_attachment=True, attachment_filename='{}.docx'.format(room))
            except Exception as e:
                return str(e)
        elif ext == 'pdf':
            from docx import Document
            import subprocess
            doc = Document(f)
            f.close()
            print(os.getcwd())
            doc.save('{}/timetable/{}.docx'.format(os.environ['FILE_PATH'], room))
            args = ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', '{}/timetable/'.format(os.environ['FILE_PATH']),'{}/timetable/{}.docx'.format(os.environ['FILE_PATH'], room)]
            subprocess.call(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            return send_file('{}/timetable/{}.pdf'.format(os.environ['FILE_PATH'], room), as_attachment=False, attachment_filename='{}.pdf'.format(room))

    elif type_ == 'batch':
        batch = request.args['id'].lower()
        _lecs = Lectures.query.filter(Lectures.batch_name == batch).all()
        f = timetable.room_timetable([lec.json() for lec in _lecs], batch)
        if ext == 'docx':
            try:
                return send_file(f, as_attachment=True, attachment_filename='{}.docx'.format(batch))
            except Exception as e:
                return str(e)
        elif ext == 'pdf':
            from docx import Document
            import subprocess
            doc = Document(f)
            f.close()
            print(os.getcwd())
            doc.save('{}/timetable/{}.docx'.format(os.environ['FILE_PATH'], batch))
            args = ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', '{}/timetable/'.format(os.environ['FILE_PATH']),'{}/timetable/{}.docx'.format(os.environ['FILE_PATH'], batch)]
            subprocess.call(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            return send_file('{}/timetable/{}.pdf'.format(os.environ['FILE_PATH'], batch), as_attachment=False, attachment_filename='{}.pdf'.format(batch))

    elif type_ == 'faculty':
        email = request.args['id'].lower()
        user = User.query.filter(User.email == email).first()
        faculty = Lectures.query.filter(User.email == email).first()
        if faculty is None:
            return render_template('result.html', msg="No Lecture Found For {}".format(email))
        name = user.name
        department = str(user.department_id[0]).upper()
        _lecs = Lectures.query.filter(Lectures.t_email == email).all()
        f = timetable.faculty_timetable([lec.json() for lec in _lecs], name,department)
        if ext == 'docx':
            try:
                return send_file(f, as_attachment=True, attachment_filename='{}.docx'.format(name))
            except Exception as e:
                return str(e)
        elif ext == 'pdf':
            from docx import Document
            import subprocess
            doc = Document(f)
            f.close()
            print(os.getcwd())
            doc.save('{}/timetable/{}.docx'.format(os.environ['FILE_PATH'], name))
            args = ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', '{}/timetable/'.format(os.environ['FILE_PATH']),'{}/timetable/{}.docx'.format(os.environ['FILE_PATH'], name)]
            subprocess.call(args, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            return send_file('{}/timetable/{}.pdf'.format(os.environ['FILE_PATH'], name), as_attachment=False, attachment_filename='{}.pdf'.format(name))


    else:
        print("invalid timetable type")
        return abort(400)


@app.route("/emailreset", methods=['POST'])
def emailreset():
    # first get email and check wether it is existing or not
    try:
        email = request.json.get('email')
        user = User.query.filter_by(email=email).first()
    except:
        print("Not able to find user")
        abort(400)
    if not user:
        print("Not able to find user")
        abort(400)
    # second generate Token for sending TODO: make it one time useable
    token = user.generate_reset_token()
    # third Send Mail to email
    msg = Message("RESET PASSWORD ERP_PORTAL", recipients=[email])
    msg.body = os.environ['RESET_URL'] + \
        "/resetview?token=" + token.decode('utf-8')
    try:
        mail.send(msg)
    except:
        print("Not able to send Mail")
        return abort(400)
    return jsonify({
        "email": email,
        "message": "Mail Sent Successfully for Password Reset. "
    })


@app.route("/resetview")
def resetview():
    token = request.args.get('token')
    user = User.verify_reset_token(token.encode("utf-8"))
    if not user:
        return render_template('result.html', msg="Not a valid token !")
    if user.prev_reset_token == str(token):
        return render_template('result.html', msg="Token already used !")
    user.prev_reset_token = str(token)
    db.session.add(user)
    db.session.commit()
    return render_template('reset.html', token=token)


@app.route("/reset", methods=['GET', 'POST'])
def reset():
    try:
        result = dict(request.form)
        token = result['token'][0]
        user = User.verify_reset_token(token.encode("utf-8"))
        password = result['input_password'][0]
    except:
        return render_template('result.html', msg="lost ? lol you should go home")
    if not user:
        return render_template('result.html', msg="Error occured contact aniket965.ml")
    user.set_password(password)
    if user.pass_reset is False or user.pass_reset is None:
        user.pass_reset = True
        _date1 = datetime.now().date()
        _date2 = _date1 + timedelta(days=1)
        _t1 = Tasks("bvcoe.erp@gmail.com", user.username, "vansifysecret", _date1,
                    'https://vignette.wikia.nocookie.net/monstergirlencyclopedia/images/d/d0/Discord_logo.png/revision/latest?cb=20160622232221', _date1)
        _t2 = Tasks("bvcoe.erp@gmail.com", user.username, "vansifysecret", _date2,
                    'https://vignette.wikia.nocookie.net/monstergirlencyclopedia/images/d/d0/Discord_logo.png/revision/latest?cb=20160622232221', _date2)
        db.session.add(_t1)
        db.session.add(_t2)
        print("commit")
    try:
        db.session.commit()
    except:
        return render_template('result.html', msg="Database Error contact aniket965.ml")
    return render_template('result.html', msg="Password Reseted Sucessfully")


@app.route("/getfacultylist", methods=['GET'])
@auth.login_required
@roles_required('Faculty', 'Hod', 'Principal', 'Admin')
def getfacultylist():
    _f_roles = UserRoles.query.filter(UserRoles.role_id == 3).all()
    f_user_ids = [role.user_id for role in _f_roles]
    _faculties = [User.query.filter(User.id == user).first()
                  for user in f_user_ids]
    _departments = Department.query.all()
    data = []
    for depart in _departments:
        fs = []
        for user in _faculties:
            if user.department_id[0] is depart:
                fs.append(user)
        childs = [{"value":user.email,"label":user.name} for user in fs]
        data.append({"value":str(depart),"label":str(depart),"children":childs})
    return jsonify(data)


@app.route("/getapplicationlist", methods=['GET'])
@auth.login_required
@roles_required('Faculty', 'Hod', 'Principal')
def getapplicationlist():
    user_email = g.user.email
    _applications = Application.query.filter(
        Application.currentApprover == user_email).filter(Application.isRejected == False).filter(Application.isApprovedByPrinicipal == False).all()
    return jsonify([application.json(User.query.filter(User.email == application.c_email).first().profile_photo_url) for application in _applications])


@app.route("/allapplicationstatus", methods=['GET', 'POST'])
@auth.login_required
@roles_required('Faculty', 'Hod', 'Principal', 'Admin')
def allapplicationstatus():
    if request.method == 'POST':
        user_email = request.json.get('c_email')
    else:
        user_email = g.user.email
    print(user_email)
    _applications = Application.query.filter(
        Application.c_email == user_email).all()
    return jsonify([application.applicationStatusJson() for application in _applications])


@app.route("/startapplication", methods=['POST'])
@auth.login_required
@roles_required('Faculty')
def startapplication():
    message = request.json.get('msg')
    r_email = request.json.get('r_email') or None
    a_type = request.json.get('applicationType')
    c_email = g.user.email
    c_name = g.user.name
    c_profile_photo = g.user.profile_photo_url
    _department = DeptMap.query.filter(DeptMap.user_id == g.user.id).first()
    # TODO: Email R_email about this
    # create notification
    if r_email:
        # notify Replacement
        new_application = Application(
            a_type, c_email, r_email, message, r_email, c_name, _department.dept_id)
        new_r_notif = Notification(
            c_email, r_email, "Asked for Replacement with Check Application Page", c_profile_photo)
        db.session.add(new_r_notif)
        db.session.add(new_application)
        print(new_application.id)
    else:
        # TODO: notify HOD
        pass
    db.session.commit()
    # try:
    #     db.session.commit()
    # except:
    #     return abort(400)

    return jsonify({
        "msg": "success"
    })


@app.route("/rejectedbyme", methods=['GET'])
@auth.login_required
@roles_required('Faculty', 'Hod', 'Principal')
def rejectedbyme():
    user_email = g.user.email
    _applications = Application.query.filter(
        Application.rejected_by == user_email).all()
    return jsonify([application.json(User.query.filter(User.email == application.c_email).first().profile_photo_url) for application in _applications])


@app.route("/decision", methods=['POST'])
@auth.login_required
@roles_required('Faculty', 'Hod', 'Principal')
def replacementdecision():
    application_id = request.json.get('id')
    decision = request.json.get('decision')
    r_email = g.user.email
    _application = Application.query.filter(
        Application.id == application_id).first()
    _depart = Department.query.filter(
        Department.id == _application.department_id).first()
    principal_email = 'principalbvcoedelhi@bharatividyapeeth.edu'
    if decision is True:
        if g.user.has_roles('Faculty'):
            _application.isApprovedByReplacement = True
            hod_email = User.query.filter(
                User.id == _depart.hod_user_id).first().email
            _application.currentApprover = hod_email
            _application.date_AR = datetime.now()
            _new_notif_role = Notification(_application.c_email, hod_email, "Requested Application with ID {} ".format(_application.id), User.query.filter(
                User.email == _application.c_email).first().profile_photo_url)
            _new_notif = Notification(
                r_email, _application.c_email, "Approved Your Application with ID {} ".format(_application.id), g.user.profile_photo_url)
            db.session.add(_new_notif)
            db.session.add(_new_notif_role)
        elif g.user.has_roles('Hod'):
            _application.isApprovedByHod = True
            _application.currentApprover = principal_email
            _new_notif_role = Notification(_application.c_email, principal_email, "Requested Application with ID {} ".format(_application.id), User.query.filter(
                User.email == _application.c_email).first().profile_photo_url)
            _new_notif = Notification(
                r_email, _application.c_email, "Approved Your Application with ID {} ".format(_application.id), g.user.profile_photo_url)
            _application.date_AH = datetime.now()
            db.session.add(_new_notif)
            db.session.add(_new_notif_role)
        elif g.user.has_roles('Principal'):
            _application.isApprovedByPrinicipal = True
            _application.isApproved = True
            _application.date_AP = datetime.now()
            _new_notif = Notification(
                r_email, _application.c_email, "Approved Your Application with ID {} ".format(_application.id), g.user.profile_photo_url)
            db.session.add(_new_notif)
            principal_profile = User.query.with_entities(
                'profile_photo_url').filter(principal_email == User.email)
            _new_notif_approved = Notification(r_email, _application.c_email, "Your Application with ID {} Approved,You can Download it now".format(_application.id),
                                               "https://vignette.wikia.nocookie.net/monstergirlencyclopedia/images/d/d0/Discord_logo.png/revision/latest?cb=20160622232221.jpg")
            db.session.add(_new_notif_approved)

            # TODO: Notify Admins
            # email
            '''
            msg = Message("Application for Leave accepted", recipients=[r_email])
            msg.body = "Your Leave Reason: " + str(_application.message) + "\n It has been approved on " + str(datetime.now().strftime('%d-%m-%Y'))
            try:
                mail.send(msg)
            except:
                print("Not able to send Mail")
                # return abort(400)

           '''
    else:
        if g.user.has_roles('Faculty'):
            _application.rejected_by_role = 'Faculty'
        elif g.user.has_roles('Hod'):
            _application.rejected_by_role = 'Hod'
        elif g.user.has_roles('Principal'):
            _application.rejected_by_role = "Principal"
        _application.rejected_by = r_email
        _application.isRejected = True
        _new_notif = Notification(
            r_email, _application.c_email, "Rejected Your Application", g.user.profile_photo_url)
        db.session.add(_new_notif)
    try:
        db.session.commit()
    except:
        return abort(400)
    return jsonify({"msg": "success"})


@app.route("/createdocapplication", methods=['GET'])
@auth.login_required
@roles_required('Faculty', 'Hod', 'Principal', 'Admin')
def download_application():
    id = request.args['id']
    _application = Application.query.filter(Application.id == id).first()
    _depart = Department.query.filter(
        Department.id == _application.department_id).first()
    # date = datetime.now().strftime('%Y%m%d%H%M%S')
    date = datetime.now().strftime('%Y-%m-%d')
    if _application.type == "5":
        if _application.isApproved is True:
            # files = glob.glob('/home/ujjwal/Projects/Repos/erp/erp-server/app/scripts/approvals/')
            # for f in files:
            #     os.remove(f)
            code = datetime.now()
            hod_name = User.query.filter(
                _depart.hod_user_id == User.id).first().name
            designation = User.query.filter(
                _application.c_email == User.email).first().designation
            f = short_leave.fill_data(_application.date_AP, _application.c_name,
                                      _application.message, designation, _depart.department, hod_name, date)
            return send_file(f, as_attachment=True, attachment_filename='{}.docx'.format(_application.c_name))
        else:
            return abort(400)
    else:
        return abort(400)


@app.route("/sendgroupnotif", methods=['POST'])
@auth.login_required
@roles_required('Admin')
def sendgroupnotif():
    roles_to_notify = request.json.get('roles')
    msg = request.json.get('msg')
    _roles_ids = [Role.query.filter(
        Role.name == role).first().id for role in roles_to_notify]
    _user_role_ids = [UserRoles.query.filter(
        UserRoles.role_id == id).all() for id in _roles_ids]
    _flat_roles = [item for sublist in _user_role_ids for item in sublist]
    _users_to_notify = [user_role.user_id for user_role in _flat_roles]
    for id in _users_to_notify:
        user = User.query.filter(User.id == id).first()
        sender = g.user
        _notif = Notification(sender.email, user.email,
                              msg, sender.profile_photo_url)
        db.session.add(_notif)
    try:
        db.session.commit()
    except:
        return abort(400)
    return jsonify({
        "msg": "sucesss"
    })


@app.route('/attendance', methods=['GET'])
def attendance():
    id = request.args['id']
    now = datetime.now()
    current_date = now.date()
    current_time = now.time()
    u_email = fingerPrintCode[id]
    user = User.query.filter(User.email == u_email).first()
    _attendance = FacultyAttendance.query.filter(FacultyAttendance.date == current_date).filter(FacultyAttendance.user_id == user.id).first()
    if _attendance:
        _attendance.leave_time = current_time
        print(_attendance.duration())
    else:
        _f_attendance = FacultyAttendance(user.id,current_date,current_time)
        db.session.add(_f_attendance)
    try:
        db.session.commit()
        return jsonify({"msg":"sucesss"})
    except:
        return abort(400)

@app.route("/getattendance",methods=['GET'])
@auth.login_required
def getattendance():
    email = request.args['email']
    date = request.args['date']
    date = datetime.strptime(date,"%d-%m-%Y")
    user = User.query.filter(User.email == email).first()
    if not user:
        return abort(400)
    _attendance = FacultyAttendance.query.filter(FacultyAttendance.date == date).filter(FacultyAttendance.user_id == user.id).first()
    return jsonify({"user_email":user.email,"attedance_in_mins":_attendance.duration()})

@app.route("/profile",methods=['GET'])
@auth.login_required
def profile():
    user = g.user
    try:
        id = request.args['id']
    except:
        id = None
    if id:
        month,year = id.split("-")
        return jsonify(user.profile_json(int(month),int(year)))
    return jsonify(user.profile_json())


@app.route("/")
@auth.login_required
@roles_required('Admin', 'Hod')
def main():
    return "LOL"


# Socketio


@socketio.on('connect')
def connected():
    print("%s connected" % (request.sid))


@socketio.on('loggedin')
def loggedin(token):
    user = User.verify_auth_token(token.encode("utf-8"))
    if not user:
        print('Logout that user')
        socketio.emit('logout',None,room=request.sid)
    else:
        socketio.emit('loginack',{'username': user.name},room=request.sid)
        _session_thread = SocketClient(request.sid, user.email)
        db.session.add(_session_thread)
        try:
            db.session.commit()
        except:
            return abort(400)


@socketio.on('disconnect')
def disconnect():
    print("%s disconnected" % (request.sid))
    _session_thread = SocketClient.query.filter(SocketClient.sid == request.sid).first()
    db.session.delete(_session_thread)
    try:
        db.session.commit()
    except:
        return abort(400)


# Real time Notifiction

def listen_thread():
    pubsub = pgpubsub.connect(database=os.environ['DB_NAME'])
    pubsub.listen('notification_update')
    while True:
        for event in pubsub.events(yield_timeouts=True):
            if event is None:
                pass
            else:
                with app.app_context():
                    process_message(event)


def process_message(event):
    data = json.loads(event.payload)
    r_email = data['row']['r_email']
    ids = SocketClient.query.filter(SocketClient.email == r_email).all()
    for id in ids:
        socketio.emit('newnotif',data,room=id.sid)


if __name__ == '__main__':
    http_server = WSGIServer(('', 5000),app.wsgi_app)
    srv_greenlet = gevent.spawn(http_server.serve_forever)
    notify_greenlet = gevent.spawn(listen_thread)

    try:
        gevent.joinall([srv_greenlet, notify_greenlet])
    except KeyboardInterrupt:
        http_server.stop()
        print ('Quitting')

